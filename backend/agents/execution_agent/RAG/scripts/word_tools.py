from __future__ import annotations

import glob
import os
from datetime import datetime
from typing import Iterable, Sequence

from docx import Document


def _get_agent_folder(subfolder: str) -> str:
    possible_paths = [
        os.path.expanduser("~/OneDrive/Desktop/agent"),
        os.path.expanduser("~/Desktop/agent"),
        os.path.expanduser("~/Documents/agent"),
    ]

    for base_path in possible_paths:
        if os.path.exists(os.path.dirname(base_path)) or os.path.exists(base_path):
            folder = os.path.join(base_path, subfolder)
            os.makedirs(folder, exist_ok=True)
            return folder

    fallback_path = os.path.join(os.path.expanduser("~"), "agent", subfolder)
    os.makedirs(fallback_path, exist_ok=True)
    return fallback_path


def _ensure_docx_extension(name: str) -> str:
    if name.lower().endswith(".docx"):
        return name
    return f"{name}.docx"


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def _print_success(path: str) -> None:
    print(f"[FILE]: {path}")
    print("EXECUTION_SUCCESS")


def _save_with_retry(doc: Document, path: str) -> str:
    try:
        doc.save(path)
        return path
    except PermissionError:
        base, ext = os.path.splitext(path)
        retry_path = f"{base}_v2{ext or '.docx'}"
        doc.save(retry_path)
        return retry_path


def _resolve_doc_path(filename: str) -> str:
    if os.path.isabs(filename) and os.path.exists(filename):
        return filename

    folder = _get_agent_folder("docs")
    search_roots = [
        os.path.expanduser("~\\Desktop"),
        os.path.expanduser("~\\OneDrive\\Desktop"),
        os.path.expanduser("~\\Documents"),
        os.path.expanduser("~\\Downloads"),
        folder,
    ]

    for root in search_roots:
        if not os.path.exists(root):
            continue
        direct = os.path.join(root, filename)
        if os.path.isfile(direct):
            return direct
        matches = glob.glob(os.path.join(root, "**", filename), recursive=True)
        if matches:
            return matches[0]

    raise FileNotFoundError(f"Could not find file: {filename}")


def doc_create(name: str | None = None) -> str:
    folder = _get_agent_folder("docs")
    filename = _ensure_docx_extension(name) if name else f"document_{_timestamp()}.docx"
    save_path = filename if os.path.isabs(filename) else os.path.join(folder, filename)

    doc = Document()
    final_path = _save_with_retry(doc, save_path)
    _print_success(final_path)
    return final_path


def doc_open(path: str) -> None:
    target = _resolve_doc_path(path)
    os.startfile(target)
    _print_success(target)


def doc_find(filename: str) -> str:
    found = _resolve_doc_path(filename)
    _print_success(found)
    return found


def doc_load(path: str) -> Document:
    resolved = _resolve_doc_path(path)
    return Document(resolved)


def doc_add_heading(doc: Document, text: str, level: int = 1) -> Document:
    doc.add_heading(text, level=level)
    return doc


def doc_add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> Document:
    if bold or italic:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        return doc

    doc.add_paragraph(text)
    return doc


def doc_add_table(doc: Document, headers: Sequence[str], rows: Iterable[Sequence[str]]) -> Document:
    table = doc.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = str(header)

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            if idx >= len(header_cells):
                break
            row_cells[idx].text = str(value)

    return doc


def doc_save(doc: Document, path: str) -> str:
    save_path = _ensure_docx_extension(path)
    if not os.path.isabs(save_path):
        save_path = os.path.join(_get_agent_folder("docs"), save_path)
    final_path = _save_with_retry(doc, save_path)
    _print_success(final_path)
    return final_path


def doc_launch() -> None:
    os.system("start winword")
    print("EXECUTION_SUCCESS")


__all__ = [
    "doc_create",
    "doc_open",
    "doc_find",
    "doc_load",
    "doc_add_heading",
    "doc_add_paragraph",
    "doc_add_table",
    "doc_save",
    "doc_launch",
]
